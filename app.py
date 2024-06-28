import os
import asyncio
from langchain_groq import ChatGroq
import requests
from typing import Type, Any
from io import BytesIO
from crewai_tools import ScrapeWebsiteTool
import streamlit as st
from streamlit import cache
from docx import Document
from pydantic.v1 import BaseModel, Field
from crewai_tools.tools.base_tool import BaseTool
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from crewai import Agent, Task, Crew
from typing import List
from langchain.output_parsers import PydanticOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_core.pydantic_v1 import BaseModel, Field, validator
# Import DuckDuckGoSearchRun: Tool for web searches via DuckDuckGo
from langchain_community.tools import DuckDuckGoSearchRun

def create_word_doc(text):
    # Create a new Document
    doc = Document()
    
    # Add text to the document
    doc.add_paragraph(text)
    
    doc.save(f'Outline.docx')

# serp_api_key = ''

# class SerpApiGoogleSearchToolSchema(BaseModel):
#     q: str = Field(..., description="Parameter defines the query you want to search. You can use anything that you would use in a regular Google search. e.g. inurl:, site:, intitle:.")
#     # tbs: str = Field("qdr:w2", description="Time filter to limit the search to the last two weeks.")

# class SerpApiGoogleSearchTool(BaseTool):
#     name: str = "Google Search"
#     description: str = "Search the internet"
#     args_schema: Type[BaseModel] = SerpApiGoogleSearchToolSchema
#     search_url: str = "https://serpapi.com/search"
    
#     def _run(
#         self,
#         q: str,
#         # tbs: str = "qdr:w2",
#         **kwargs: Any,
#     ) -> Any:
#         global serp_api_key
#         payload = {
#             "engine": "google",
#             "q": q,
#             # "tbs": tbs,
#             "api_key": serp_api_key,
#         }
#         headers = {
#             'content-type': 'application/json'
#         }
    
#         response = requests.get(self.search_url, headers=headers, params=payload)
#         results = response.json()
    
#         summary = ""
#         for key in ['answer_box_list', 'answer_box', 'organic_results', 'sports_results', 'knowledge_graph', 'top_stories']:
#             if key in results:
#                 summary += str(results[key])
#                 break
        
#         print(summary)
        
#         return summary
   
# def generate_text(llm, chapter, sub_topics, serpapi_key):
#     inputs = {'chapter': chapter, 'sub_topics':sub_topics}
#     search_tool = SerpApiGoogleSearchTool()
    
#     scrape_tool = ScrapeWebsiteTool(
#         name="website_scraper",
#         description="""Scrape content from web pages. Action Input should look like this:
#                        {"website_url": "<URL of the webpage to scrape>"}""",
#     )

#     researcher_agent = Agent(
#         role='Book Chapter Researcher',
#         goal='Search for really informative sources on the given chapter topic "{chapter}" and sub topics {sub_topics}, find unique URLs containing detailed information, and scrape relevant information from these URLs.',
#         backstory=(
#             "An experienced researcher with strong skills in web scraping, fact-finding, and "
#             "analyzing in-depth information to provide comprehensive material for writing book chapters."
#         ),
#         verbose=True,
#         allow_delegation=False,
#         max_iter = 5,
#         llm=llm
#     )

#     writer_agent = Agent(
#         role='Book Chapter Writer',
#         goal='Write a detailed, engaging, and informative book chapter based on the information found by the researcher using the format specified.',
#         backstory=("An experienced writer with a background in book writing and content creation. "
#                    "Skilled in crafting compelling narratives and distilling complex information into "
#                    "engaging content for book chapters."),
#         verbose=True,
#         allow_delegation=False,
#         max_iter = 5,
#         llm=llm
#     )

#     reviewer_agent = Agent(
#         role='Content Reviewer',
#         goal='Review and refine content drafts to ensure they meet high standards of quality and coherence like professional book chapters. chapter name is {chapter}, also consider the sub topics {sub_topics} while writing the chapter',
#         backstory=("A meticulous reviewer with extensive experience in editing and proofreading, "
#                    "known for their keen eye for detail and commitment to maintaining the highest quality standards in published content."),
#         verbose=True,
#         allow_delegation=False,
#         llm=llm
#     )
    
#     final_writer_agent = Agent(
#         role='Final Content Writer',
#         goal='Compile, refine, and structure all reviewed and approved content into a cohesive and engaging book chapter format. Ensure that the final product is polished, logically structured, and ready for publication, providing a seamless and informative reading experience for the audience.',
#         backstory=("An accomplished writer and editor with extensive experience in book writing, content creation, and editorial management. "
#                    "Known for their ability to craft compelling narratives and ensure consistency and quality across all sections of a publication. "
#                    "With a keen eye for detail and a deep understanding of audience engagement, this writer excels in transforming raw content into polished, professional-grade book chapters that captivate readers and deliver clear, valuable insights."),
#         verbose=True,
#         allow_delegation=False,
#         llm=llm
#     )
    
#     task_researcher = Task(
#         description=(f'Research and identify the most informative sources on the topic of {chapter} and sub topics {sub_topics}'
#                      'Scrape detailed content from relevant websites to gather comprehensive material.'),
#         agent=researcher_agent,
#         expected_output=('A list of informative sources with their respective website URLs. '
#                          'Scraped content from all URLs that can be used further by the writer.'),
#         tools=[search_tool, scrape_tool]
#     )

#     task_writer = Task(
#         description=('Write a detailed book chapter based on the information found by the researcher. '
#                      'Ensure the chapter is informative, engaging, and provides clear insights into the topic.'),
#         agent=writer_agent,
#         expected_output=('A drafted book chapter with comprehensive information, structured in a clear and engaging manner. '
#                          'Include citations and references to the original sources found by the researcher.')
#     )

#     task_reviewer = Task(
#         description=('Review the drafted book chapter provided by the writer for accuracy, coherence, and quality. '
#                      'Ensure that the content is free from errors and meets the publication standards.'),
#         agent=reviewer_agent,
#         expected_output=('Reviewed book chapter with suggestions for improvements, if any. '
#                          'Final versions of the chapter that are ready for inclusion in the book.')
#     )

#     task_final_writer = Task(
#         description=('Compile the reviewed and refined content into a well-structured book chapter format. '
#                      'Ensure the chapter is logically organized and flows seamlessly.'),
#         agent=final_writer_agent,
#         expected_output=(
#             """Final book chapter document with all the reviewed content, formatted and ready for publication. 
#             The chapter should include:
#             - An engaging introduction to the topic. Start with an engaging hook: a question, a quote, or a surprising fact.
#             - Main content sections:
#                 - Detailed subsections covering different aspects of the topic.
#                 - Each subsection should have:
#                     - A small introduction.
#                     - Detailed information presented in a clear and logical manner.
#                     - Explanation of the importance or relevance of the information.
#             - Conclusion:
#                 - Summarize the chapter by wrapping up the main points and providing a final thought or conclusion.
#             """
#         )
#     )


#     crew = Crew(
#         agents=[researcher_agent, writer_agent, reviewer_agent, final_writer_agent],
#         tasks=[task_researcher, task_writer, task_reviewer, task_final_writer],
#         verbose=2,
#         context={"Book Chapter Topic is ": chapter}
#     )

#     result = crew.kickoff(inputs=inputs)

#     return result

def generate_text(llm, chapter, sub_topics):
    inputs = {'chapter': chapter, 'sub_topics':sub_topics}

    # Initialize DuckDuckGo web search tool: Enables real-time fact-finding for debates
    search_tool = DuckDuckGoSearchRun(
        name="duckduckgo_search",
        description="""Search the web using DuckDuckGo. Action Input should look like this:
                       {"query": "<Whatever you want to search>"}"""
    )

    # Define Book chapter Researcher Agent
    book_chapter_researcher = Agent(
        role='Book chapter Content Researcher',
        goal='Conduct thorough research to uncover compelling insights for engaging Book chapter and sub_topics content.',
        backstory=("An experienced content strategist with a knack for analyzing trends and audience behavior, "
                   "delivering actionable insights for high-quality Book chapter content."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    # Define Book chapter Writer Agent
    book_chapter_writer = Agent(
        role='Book chapter Writer',
        goal=f'Craft authoritative and engaging Book chapter {chapter} content and sub topics {sub_topics} content that resonates with the audience and establishes the brand as a leader.',
        backstory=("A seasoned writer known for distilling complex topics into captivating stories, with a deep understanding of audience psychology."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    # Define Book chapter Reviewer Agent
    book_chapter_reviewer = Agent(
        role='Content Reviewer',
        goal='Review and refine Book chapter drafts to ensure they meet high standards of quality and impact.',
        backstory=("An expert editor with a meticulous eye for detail, known for elevating content to publication-ready standards."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    # Define Task for Researcher
    task_researcher = Task(
        description=(f"Research the latest trends and insights on {chapter}. Identify key developments, emerging trends, unique perspectives, and content ideas."),
        agent=book_chapter_researcher,
        expected_output=(
            f"1. Overview and background of {chapter} and sub topics {sub_topics}.\n"
            "2. Recent key developments.\n"
            "3. Emerging trends and innovative approaches.\n"
            "4. Unique angles and untapped opportunities.\n"
            "5. Potential content ideas with brief descriptions.\n"
            "6. List of relevant sources."
        ),
        tools=[search_tool]
    )

    # Define Task for Writer
    task_writer = Task(
        description=(f"Based on the research report, craft an engaging and authoritative Book chapter post on {chapter}."),
        agent=book_chapter_writer,
        expected_output=(
            "1. Engaging introduction with a hook.\n"
            "2. Use of deatiled exploration of key developments.\n"
            "3. Use of emerging trends and innovative ideas in content.\n"
            "4. Use of unique angles and perspectives in content.\n"
            "5. Clear explanations of complex concepts.\n"
            "7. Compelling conclusion.\n"
        )
    )

    # Define Task for Reviewer
    task_reviewer = Task(
        description=(f"Review the drafted Book chapter content on {chapter}, providing detailed feedback and revisions for quality and impact."),
        agent=book_chapter_reviewer,
        expected_output=(
            "1. Overall content assessment.\n"
            "2. Identification of inaccuracies and gaps.\n"
            "3. Suggestions for improving flow and readability.\n"
            "4. Recommendations for tone and voice.\n"
            "5. Edits for grammar and punctuation.\n"
            "6. Final assessment of readiness."
        )
    )

    # Define Task for Final Writer
    task_final_writer = Task(
        description=(f"Revise the Book chapter content on {chapter} and sub topic {sub_topics} based on the reviewer's feedback, ensuring it meets high standards."),
        agent=book_chapter_writer,
        expected_output=(
            "1. Factually accurate and corrected content.\n"
            "2. Clear, well-structured flow.\n"
            "3. Concise and engaging language.\n"
            "4. Consistent tone and voice.\n"
            "5. Enhanced insights.\n"
            "6. Addressed reviewer feedback.\n"
            "7. Creative and engaging Book chapter title.\n"
            "8. Final draft of at least 1000 words."
        )
    )

    # Initialize Crew: Coordinates agents and tasks for structured Book chapter content workflow
    crew = Crew(
        agents=[book_chapter_researcher, book_chapter_writer, book_chapter_reviewer, book_chapter_writer],
        tasks=[task_researcher, task_writer, task_reviewer, task_final_writer],
        verbose=2,
        context={"Chapter Topic is ": chapter}
    )

    # Start the workflow and generate the result
    result = crew.kickoff(inputs=inputs)

    return result

def book_overview(topic,llm):
    class overview(BaseModel):
        outline: List[dict] = Field(description="chapter and points to be included in each chapter")
        title: str = Field(description="compelling book title based on the topic")

    # And a query intented to prompt a language model to populate the data structure.
    prompt_topic = f"You are a professional author. You need to Generate a detailed outline for a book on {topic}. the outline should have interesting chapter name and the points ot be included in the chapter."

    # Set up a parser + inject instructions into the prompt template.
    parser = PydanticOutputParser(pydantic_object=overview)

    prompt = PromptTemplate(
        template="Answer the user query.\n{format_instructions}\n{query}\n",
        input_variables=["query"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = prompt | llm | parser

    output = chain.invoke({"query": prompt_topic})
    return output


def main():
    
    st.header('AI E book Generator')
    mod = None
    
    global serp_api_key
    
    with st.sidebar:
        with st.form('Gemini/OpenAI/Groq'):
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI','Groq'))
            api_key = st.text_input(f'Enter your API key', type="password")
            # serp_api_key = st.text_input(f'Enter your SerpAPI key', type="password")
            submitted = st.form_submit_button("Submit")

    if api_key and serp_api_key:
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                os.environ["OPENAI_API_KEY"] = api_key
                llm = ChatOpenAI(temperature=0.6, max_tokens=2000)
                print("OpenAI Configured")
                return llm

            llm = asyncio.run(setup_OpenAI())
            mod = 'Gemini'


        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key
                )
                print("Gemini Configured")
                return llm

            llm = asyncio.run(setup_gemini())
            mod = 'Gemini'

            
        elif model == 'Groq':
            async def setup_groq():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGroq(
                    api_key = api_key,
                    model = 'llama3-70b-8192'
                )
                return llm

            llm = asyncio.run(setup_groq())
            mod = 'Groq'

        

        topic = st.text_input("Enter the E book topic:")

        if st.button("Generate E book Outline:"):
            with st.spinner("Generating content..."):
                output = book_overview(topic,llm)
                outline_str = ""
                cnt=1
                for i in output.outline:
                    outline_str = outline_str + (f"Chapter {cnt}: {i['chapter']}") + '\n'
                    cnt_p = 1
                    cnt = cnt +1
                    for j in i['points']:
                        outline_str = outline_str + (f"{cnt_p}. {j}") + '\n'
                        cnt_p = cnt_p +1
                    outline_str = outline_str + '\n'

                # print(outline_str)

                # create_word_doc(outline_str)

                list_dict = output.outline

                if "main_variable" not in st.session_state:
                        st.session_state.main_variable = list_dict
                        st.session_state.llm = llm
                        st.session_state.output = output
                        # st.session_state.serp_api_key = serp_api_key
                        st.session_state.outline_str = outline_str


                st.header(output.title)
                # cnt = 0
                # for i in output.outline:
                #     keys = i.keys()
                #     key_list = list(keys)
                #     chapter = i[key_list[0]]
                #     sub_topics = i[key_list[1]]

                #     st.write(chapter)
                #     st.write(sub_topics)
                st.write(outline_str)
                
                st.write("Are you satisfied with the outline?")

                # if st.button("Generate E Book"):
                #     st.switch_page("pages/Generate.py")
                # if st.button("Edit"):
                #     st.switch_page("pages/edit_form.py")

        satisfaction = st.selectbox(
            "Are you satisfied with the outline?",
            ["Select an option", "Yes", "No"]
        )

        if satisfaction == "Yes":
            if "main_variable" not in st.session_state:
                st.session_state.main_variable = list_dict
            st.switch_page("pages/generate.py")
            # for i in output.outline:
            #     keys = i.keys()
            #     key_list = list(keys)
            #     chapter = i[key_list[0]]
            #     sub_topics = i[key_list[1]]
            #     generated_content = generate_text(llm, chapter, sub_topics, serp_api_key)

            #     word_file = create_word_doc(generated_content,cnt)
            #     cnt = cnt + 1

        if satisfaction == "No":
            if "main_variable" not in st.session_state:
                st.session_state.main_variable = list_dict
            st.switch_page("pages/edit_form.py")

                    


    #             content_lines = generated_content.split('\n')
    #             first_line = content_lines[0]
    #             remaining_content = '\n'.join(content_lines[1:])

    #             st.markdown(first_line)
    #             st.markdown(remaining_content)

    #             doc = Document()

    #             doc.add_heading(topic, 0)
    #             doc.add_paragraph(first_line)
    #             doc.add_paragraph(remaining_content)

    #             buffer = BytesIO()
    #             doc.save(buffer)
    #             buffer.seek(0)
    
    
    #             st.download_button(
    #     label="Download as Word Document",
    #     data=buffer,
    #     file_name=f"{topic}.docx",
    #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # )

if __name__ == "__main__":
    main()
